from __future__ import annotations

from datetime import date
from pathlib import Path
import shutil

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_PATH = Path("docs/btc_project_report.docx")
LEGACY_OUTPUT_PATH = Path("btc_project_report.docx")
DATASET_PATH = Path("processed/final_dataset_v2_plus.csv")


def set_run_font(run, size: float = 12, bold: bool = False, font_name: str = "Times New Roman") -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:eastAsia"), "SimSun")


def set_run_east_asia_font(run, font_name: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_line_spacing(paragraph, line_spacing: float = 1.5) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_after = Pt(6)


def add_heading(document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    set_run_east_asia_font(run, "SimHei")
    set_paragraph_line_spacing(paragraph, 1.2)


def add_paragraph(document, text: str, indent: bool = True) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    run = paragraph.add_run(text)
    set_run_font(run)
    set_paragraph_line_spacing(paragraph)


def add_bullet(document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run)
    set_paragraph_line_spacing(paragraph, 1.3)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=11, bold=bold)
    set_paragraph_line_spacing(paragraph, 1.2)


def add_table(document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=10)


def configure_document(document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    footer = section.footer
    footer_para = footer.paragraphs[0]
    add_page_number(footer_para)


def read_dataset_summary(path: Path) -> dict[str, object]:
    frame = pd.read_csv(path)
    return {
        "rows": len(frame),
        "columns": len(frame.columns),
        "date_min": str(frame["date"].min()),
        "date_max": str(frame["date"].max()),
    }


def build_report():
    summary = read_dataset_summary(DATASET_PATH)

    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("BTC 数据工程项目交付报告")
    set_run_font(title_run, size=18, bold=True)
    set_run_east_asia_font(title_run, "SimHei")
    set_paragraph_line_spacing(title, 1.2)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("——当前正式 BTC 日频数据集说明")
    set_run_font(subtitle_run, size=14)
    set_run_east_asia_font(subtitle_run, "KaiTi")
    set_paragraph_line_spacing(subtitle, 1.2)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"生成日期：{date.today().isoformat()}\n当前正式数据集：processed/final_dataset_v2_plus.csv\n"
    )
    set_run_font(meta_run, size=11)
    set_paragraph_line_spacing(meta, 1.2)

    add_heading(document, "摘 要", 1)
    add_paragraph(
        document,
        f"本报告用于说明当前正式 BTC 日频数据集的构成、来源、构建流程、验证情况与交付方式。当前最终训练表为 `processed/final_dataset_v2_plus.csv`，共有 {summary['rows']} 行、{summary['columns']} 列，日期范围为 {summary['date_min']} 至 {summary['date_max']}。该数据集以 Binance 为价格主源，以 Blockchain.com 提供链上活跃度与网络难度信息，以 FRED 提供宏观变量，以 Alternative.me 提供情绪特征，并在最终输出前统一完成日期对齐、特征构造、前向填充、1 期滞后与缺失样本过滤。因此，这份数据集已经是一份可直接供 baseline、RNN 和 SHAP 解释使用的正式输入数据包。",
    )
    add_paragraph(document, "关键词：BTC；日频数据集；特征工程；情绪特征；网络难度；防数据泄露", indent=False)

    add_heading(document, "1 当前数据集概况", 1)
    add_paragraph(
        document,
        "当前数据集面向 BTC 单资产方向预测任务构建，目标变量定义为 `target_t = 1[close_t > close_{t-1}]`。数据集采用日频粒度，以统一的 `date` 作为主键，并保证最终进入建模的全部特征都只来自预测时点之前可获得的信息。这样做的目的是让后续的 baseline、RNN 与解释分析都建立在一致且可复现的输入基础之上。",
    )

    add_heading(document, "2 当前数据来源", 1)
    add_table(
        document,
        ["来源", "当前用途", "主要字段"],
        [
            ["Binance Spot API", "BTC 价格主表", "open, high, low, close, volume, quote_volume, trade_count"],
            ["Blockchain.com Charts API", "链上活跃度与网络难度", "active_addresses, tx_count, hash_rate, fees, difficulty"],
            ["FRED", "宏观变量", "vix, sp500, dxy_proxy"],
            ["Alternative.me", "情绪特征", "fear_greed_index"],
        ],
    )
    add_paragraph(
        document,
        "这些数据源共同构成当前数据集的正式输入基础。价格主源提供市场交易信息，链上源补充网络状态，宏观源描述外部金融环境，情绪源提供加密市场情绪刻画，使最终训练表不仅包含价格信息，也包含市场状态与网络状态信号。",
    )

    add_heading(document, "3 当前字段结构", 1)
    add_bullet(document, "价格列：open, high, low, close, volume, quote_volume, trade_count")
    add_bullet(document, "链上列：active_addresses, tx_count, hash_rate, fees, difficulty")
    add_bullet(document, "宏观列：vix, sp500_return, dxy_proxy")
    add_bullet(document, "情绪列：fear_greed_index")
    add_bullet(document, "衍生列：ret_1d, ret_3d, volatility_7d, ma_ratio_7_30, high_low_spread, volume_change, active_addresses_change, fees_per_tx, difficulty_change")
    add_bullet(document, "标签列：target")

    add_heading(document, "4 当前构建流程", 1)
    add_paragraph(
        document,
        "当前正式构建流程由一组独立脚本组成：先抓取 Binance 价格原始表，再抓取 Blockchain.com 链上原始表、FRED 宏观原始表、Alternative.me 情绪原始表和 Blockchain.com 难度原始表，最后由 `notebooks/build_dataset_v2_plus.py` 完成统一构建。构建阶段会先标准化日期和数值类型，再构造价格衍生特征与链上衍生特征，对外部源做前向填充，并在最终表中对全部建模特征统一执行 1 期滞后。",
    )
    add_paragraph(
        document,
        "与此同时，最新 Binance 日期会在最终导出前被剔除，以避免未完成日线进入训练样本。经过这些步骤后，正式输出表中的保留特征列和标签列不存在缺失值，可以直接进入后续建模阶段。",
    )

    add_heading(document, "5 当前验证结果", 1)
    add_table(
        document,
        ["检查项", "当前结果"],
        [
            ["最终数据集", f"{summary['rows']} 行 / {summary['columns']} 列"],
            ["日期范围", f"{summary['date_min']} -> {summary['date_max']}"],
            ["关键字段", "fear_greed_index, difficulty, difficulty_change"],
            ["当前验证重点", "最新日线剔除、target 稳定、统一滞后、runner 顺序正确"],
        ],
    )
    add_paragraph(
        document,
        "当前正式仓库通过回归测试保护核心语义，包括：最新未完成 Binance 日线不会进入最终训练表，`target` 构造逻辑稳定，外部特征和衍生特征统一滞后 1 期，正式下游 runner 只执行当前保留流程。",
    )

    add_heading(document, "6 当前交付物", 1)
    add_bullet(document, "主数据集：`processed/final_dataset_v2_plus.csv`")
    add_bullet(document, "字段说明：`docs/data_dictionary_v2_plus.md`")
    add_bullet(document, "QA 摘要：`docs/qa_report_v2_plus.md`")
    add_bullet(document, "数据源说明：`docs/source_note_v2_plus.md`")
    add_bullet(document, "正式报告：`docs/btc_project_report.docx`")

    add_heading(document, "7 使用建议", 1)
    add_bullet(document, "baseline、RNN 和 SHAP 解释应统一使用 `processed/final_dataset_v2_plus.csv` 作为正式输入。")
    add_bullet(document, "建模时应延续当前数据集的统一滞后语义，不要在下游重新定义时间可得性规则。")
    add_bullet(document, "字段含义、来源和 QA 结果应分别以数据字典、source note 和 QA 报告为准。")

    add_heading(document, "8 结论", 1)
    add_paragraph(
        document,
        "当前仓库已经形成一套完整的 BTC 日频正式数据集交付包。它不仅包含最终训练表本身，还包含字段说明、质量摘要、数据源说明和正式报告，能够为后续建模、解释和汇报工作提供统一、清晰且可复现的数据基础。",
    )

    return document


def main() -> None:
    document = build_report()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT_PATH))
    if OUTPUT_PATH.exists():
        shutil.copyfile(OUTPUT_PATH, LEGACY_OUTPUT_PATH)
    print(f"Saved report to {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
